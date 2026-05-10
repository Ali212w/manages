# app/services/document_processor.py

import os
import PyPDF2
import docx
import openpyxl
import pandas as pd
from PIL import Image
# import pytesseract
import csv
import json
from datetime import datetime
from werkzeug.utils import secure_filename
from flask import current_app
import langdetect
from langdetect import detect
import re

class DocumentProcessor:
    """معالجة المستندات واستخراج النصوص"""
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'word': ['.docx', '.doc'],
        'excel': ['.xlsx', '.xls'],
        'csv': ['.csv'],
        'text': ['.txt', '.md', '.rst'],
        'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp']
    }
    
    def __init__(self):
        self.supported_formats = self.SUPPORTED_FORMATS
    
    def process_file(self, file, file_path=None):
        """معالجة ملف واستخراج محتواه"""
        result = {
            'success': False,
            'text': '',
            'text_ar': '',
            'text_en': '',
            'language': 'unknown',
            'metadata': {},
            'tables': [],
            'error': None
        }
        
        try:
            filename = file.filename if hasattr(file, 'filename') else os.path.basename(file_path)
            ext = os.path.splitext(filename)[1].lower()
            
            # تحديد نوع الملف
            file_type = self._get_file_type(ext)
            
            # استخراج النص حسب النوع
            if file_type == 'pdf':
                text, metadata = self._extract_from_pdf(file, file_path)
                result['text'] = text
                result['metadata'] = metadata
                
            elif file_type == 'word':
                text = self._extract_from_word(file, file_path)
                result['text'] = text
                
            elif file_type == 'excel':
                text, tables = self._extract_from_excel(file, file_path)
                result['text'] = text
                result['tables'] = tables
                
            elif file_type == 'csv':
                text, tables = self._extract_from_csv(file, file_path)
                result['text'] = text
                result['tables'] = tables
                
            elif file_type == 'image':
                text = self._extract_from_image(file, file_path)
                result['text'] = text
                
            elif file_type == 'text':
                text = self._extract_from_text(file, file_path)
                result['text'] = text
            
            # تحديد لغة النص
            if result['text']:
                result['language'] = self._detect_language(result['text'])
                
                # فصل النص العربي والإنجليزي
                result['text_ar'], result['text_en'] = self._separate_languages(result['text'])
            
            result['success'] = True
            
        except Exception as e:
            result['error'] = str(e)
            current_app.logger.error(f"Error processing file {filename}: {str(e)}")
        
        return result
    
    def _get_file_type(self, ext):
        """تحديد نوع الملف من الامتداد"""
        for file_type, extensions in self.SUPPORTED_FORMATS.items():
            if ext in extensions:
                return file_type
        return 'unknown'
    
    def _extract_from_pdf(self, file, file_path):
        """استخراج النص من PDF"""
        text = ""
        metadata = {}
        
        try:
            if file_path:
                pdf_file = open(file_path, 'rb')
            else:
                # حفظ الملف مؤقتاً
                temp_path = f"/tmp/{secure_filename(file.filename)}"
                file.save(temp_path)
                pdf_file = open(temp_path, 'rb')
            
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # استخراج الميتاداتا
            if pdf_reader.metadata:
                metadata = {
                    'title': pdf_reader.metadata.get('/Title', ''),
                    'author': pdf_reader.metadata.get('/Author', ''),
                    'subject': pdf_reader.metadata.get('/Subject', ''),
                    'creator': pdf_reader.metadata.get('/Creator', ''),
                    'producer': pdf_reader.metadata.get('/Producer', ''),
                    'page_count': len(pdf_reader.pages)
                }
            
            # استخراج النص من كل صفحة
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n--- صفحة {page_num + 1} ---\n"
                    text += page_text + "\n"
            
            pdf_file.close()
            
            # تنظيف الملف المؤقت
            if not file_path and os.path.exists(temp_path):
                os.remove(temp_path)
                
        except Exception as e:
            current_app.logger.error(f"PDF extraction error: {str(e)}")
        
        return text, metadata
    
    def _extract_from_word(self, file, file_path):
        """استخراج النص من Word"""
        text = ""
        
        try:
            if file_path:
                doc = docx.Document(file_path)
            else:
                temp_path = f"/tmp/{secure_filename(file.filename)}"
                file.save(temp_path)
                doc = docx.Document(temp_path)
            
            # استخراج النص من الفقرات
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + "\n"
            
            # استخراج النص من الجداول
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            # تنظيف الملف المؤقت
            if not file_path and os.path.exists(temp_path):
                os.remove(temp_path)
                
        except Exception as e:
            current_app.logger.error(f"Word extraction error: {str(e)}")
        
        return text
    
    def _extract_from_excel(self, file, file_path):
        """استخراج النص من Excel"""
        text = ""
        tables = []
        
        try:
            if file_path:
                df_dict = pd.read_excel(file_path, sheet_name=None)
            else:
                df_dict = pd.read_excel(file, sheet_name=None)
            
            for sheet_name, df in df_dict.items():
                text += f"\n--- ورقة: {sheet_name} ---\n"
                
                # تحويل DataFrame إلى نص
                table_data = df.to_string()
                text += table_data + "\n"
                
                # حفظ الجدول كـ JSON
                tables.append({
                    'sheet': sheet_name,
                    'columns': df.columns.tolist(),
                    'data': df.to_dict(orient='records')
                })
                
        except Exception as e:
            current_app.logger.error(f"Excel extraction error: {str(e)}")
        
        return text, tables
    
    def _extract_from_csv(self, file, file_path):
        """استخراج النص من CSV"""
        text = ""
        tables = []
        
        try:
            if file_path:
                df = pd.read_csv(file_path)
            else:
                df = pd.read_csv(file)
            
            text = df.to_string()
            
            tables.append({
                'sheet': 'data',
                'columns': df.columns.tolist(),
                'data': df.to_dict(orient='records')
            })
            
        except Exception as e:
            current_app.logger.error(f"CSV extraction error: {str(e)}")
        
        return text, tables
    
    def _extract_from_image(self, file, file_path):
        """استخراج النص من الصور باستخدام OCR"""
        text = ""
        
        try:
            if file_path:
                image = Image.open(file_path)
            else:
                image = Image.open(file)
            
            # محاولة OCR باللغتين العربية والإنجليزية
            text = pytesseract.image_to_string(image, lang='ara+eng')
            
        except Exception as e:
            current_app.logger.error(f"Image OCR error: {str(e)}")
        
        return text
    
    def _extract_from_text(self, file, file_path):
        """استخراج النص من الملفات النصية"""
        text = ""
        
        try:
            if file_path:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                text = file.read().decode('utf-8')
                
        except UnicodeDecodeError:
            # محاولة بترميز مختلف
            try:
                if file_path:
                    with open(file_path, 'r', encoding='cp1256') as f:
                        text = f.read()
                else:
                    file.seek(0)
                    text = file.read().decode('cp1256')
            except:
                pass
                
        except Exception as e:
            current_app.logger.error(f"Text extraction error: {str(e)}")
        
        return text
    
    def _detect_language(self, text):
        """تحديد لغة النص"""
        try:
            # خذ أول 500 حرف للكشف
            sample = text[:500].strip()
            if not sample:
                return 'unknown'
            
            lang = detect(sample)
            
            # تعيين اللغة
            if lang == 'ar':
                return 'arabic'
            elif lang == 'en':
                return 'english'
            else:
                return lang
                
        except:
            # فحص بسيط للأحرف العربية
            arabic_chars = re.findall(r'[\u0600-\u06FF]', text[:1000])
            if len(arabic_chars) > 10:
                return 'arabic'
            return 'unknown'
    
    def _separate_languages(self, text):
        """فصل النص العربي عن الإنجليزي"""
        arabic_text = ""
        english_text = ""
        
        # أنماط للغات
        arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]+')
        english_pattern = re.compile(r'[a-zA-Z]+')
        
        lines = text.split('\n')
        for line in lines:
            # كلمات عربية
            arabic_words = arabic_pattern.findall(line)
            if arabic_words:
                arabic_text += ' '.join(arabic_words) + ' '
            
            # كلمات إنجليزية
            english_words = english_pattern.findall(line)
            if english_words:
                english_text += ' '.join(english_words) + ' '
        
        return arabic_text.strip(), english_text.strip()