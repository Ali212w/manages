"""
template_generator.py - نظام توليد قوالب المشاريع بتنسيقات متعددة
"""
import os
import json
import csv
from datetime import datetime
from flask import send_file, current_app
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import arabic_reshaper
from bidi.algorithm import get_display
import io

class TemplateGenerator:
    """مولد قوالب المشاريع بتنسيقات متعددة مع دعم العربية"""
    
    def __init__(self):
        self.supported_formats = ['excel', 'word', 'pdf', 'csv', 'json']
        
        # دعم اللغة العربية للـ PDF
        try:
            # تسجيل خط يدعم العربية
            pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
            pdfmetrics.registerFont(TTFont('ArialBD', 'arialbd.ttf'))
        except:
            pass
    
    def _prepare_arabic_text(self, text):
        """تجهيز النص العربي للعرض"""
        if not text:
            return ''
        try:
            reshaped_text = arabic_reshaper.reshape(text)
            return get_display(reshaped_text)
        except:
            return text
    
    # ============================================
    # توليد قالب Excel
    # ============================================
    
    def generate_excel_template(self, project_data=None):
        """توليد قالب Excel مع دعم العربية"""
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # صفحة معلومات المشروع
            project_df = pd.DataFrame({
                'الحقل': [
                    'اسم المشروع',
                    'رمز المشروع',
                    'وصف المشروع',
                    'تاريخ البدء (YYYY-MM-DD)',
                    'تاريخ الانتهاء (YYYY-MM-DD)',
                    'الميزانية',
                    'العملة',
                    'موقع المشروع',
                    'العميل',
                    'مدير المشروع',
                    'نوع المشروع'
                ],
                'القيمة': [
                    project_data.get('name', '') if project_data else '',
                    project_data.get('code', '') if project_data else '',
                    project_data.get('description', '') if project_data else '',
                    project_data.get('start_date', '') if project_data else '',
                    project_data.get('end_date', '') if project_data else '',
                    project_data.get('budget', '') if project_data else '',
                    project_data.get('currency', 'SAR') if project_data else 'SAR',
                    project_data.get('location', '') if project_data else '',
                    project_data.get('client', '') if project_data else '',
                    project_data.get('manager', '') if project_data else '',
                    project_data.get('type', '') if project_data else ''
                ],
                'ملاحظات': [
                    'اسم المشروع بالعربية أو الإنجليزية',
                    'رمز فريد للمشروع',
                    'وصف مختصر للمشروع',
                    'صيغة التاريخ: 2024-01-15',
                    'صيغة التاريخ: 2024-12-31',
                    'القيمة بالأرقام فقط',
                    'SAR, USD, EUR, إلخ',
                    'عنوان أو موقع المشروع',
                    'اسم العميل',
                    'اسم مدير المشروع',
                    'هندسي، برمجي، إنشائي، إلخ'
                ]
            })
            project_df.to_excel(writer, sheet_name='معلومات المشروع', index=False)
            
            # تنسيق الخلايا
            worksheet = writer.sheets['معلومات المشروع']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
            
            # صفحة المهام
            tasks_data = []
            if project_data and project_data.get('tasks'):
                for i, task in enumerate(project_data['tasks'], 1):
                    tasks_data.append({
                        'الرقم': i,
                        'اسم المهمة': task.get('name', ''),
                        'الوصف': task.get('description', ''),
                        'المسؤول': task.get('assigned_to', ''),
                        'تاريخ البدء': task.get('start_date', ''),
                        'تاريخ الانتهاء': task.get('end_date', ''),
                        'المدة (أيام)': task.get('duration', ''),
                        'الأولوية': task.get('priority', ''),
                        'المهام السابقة': task.get('depends_on', ''),
                        'الموارد': task.get('resources', '')
                    })
            else:
                # قالب فارغ مع 10 صفوف
                for i in range(1, 11):
                    tasks_data.append({
                        'الرقم': i,
                        'اسم المهمة': '',
                        'الوصف': '',
                        'المسؤول': '',
                        'تاريخ البدء': '',
                        'تاريخ الانتهاء': '',
                        'المدة (أيام)': '',
                        'الأولوية': 'متوسطة',
                        'المهام السابقة': '',
                        'الموارد': ''
                    })
            
            tasks_df = pd.DataFrame(tasks_data)
            tasks_df.to_excel(writer, sheet_name='المهام', index=False)
            
            # تنسيق صفحة المهام
            worksheet = writer.sheets['المهام']
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 40)
                worksheet.column_dimensions[column_letter].width = adjusted_width
            
            # صفحة التعليمات
            instructions_df = pd.DataFrame({
                'العمود': [
                    'اسم المهمة',
                    'المدة (أيام)',
                    'الأولوية',
                    'المهام السابقة',
                    'الموارد'
                ],
                'الوصف': [
                    'اسم المهمة (عربي أو إنجليزي)',
                    'عدد الأيام المتوقعة لإنجاز المهمة',
                    'عالية - متوسطة - منخفضة',
                    'أرقام المهام التي تعتمد عليها هذه المهمة (مفصولة بفواصل)',
                    'المواد أو المعدات المطلوبة (مفصولة بفواصل)'
                ],
                'مثال': [
                    'صب الخرسانة',
                    '5',
                    'عالية',
                    '1,2',
                    'خرسانة، حديد، مضخة خرسانة'
                ]
            })
            instructions_df.to_excel(writer, sheet_name='تعليمات', index=False)
        
        output.seek(0)
        return output
    
    # ============================================
    # توليد قالب Word
    # ============================================
    
    def generate_word_template(self, project_data=None):
        """توليد قالب Word مع دعم العربية"""
        doc = Document()
        
        # إضافة خط يدعم العربية
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        
        # عنوان المستند
        title = doc.add_heading('نموذج مشروع', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(24)
        title.runs[0].font.bold = True
        
        doc.add_paragraph('_' * 50)
        
        # ============================================
        # معلومات المشروع
        # ============================================
        doc.add_heading('معلومات المشروع', level=1)
        
        project_fields = [
            ('اسم المشروع', 'project_name', 'مشروع إنشاء مبنى إداري'),
            ('رمز المشروع', 'project_code', 'PRJ-2024-001'),
            ('وصف المشروع', 'description', 'وصف مختصر للمشروع'),
            ('تاريخ البدء', 'start_date', '2024-01-15'),
            ('تاريخ الانتهاء', 'end_date', '2024-12-31'),
            ('الميزانية', 'budget', '1,500,000'),
            ('العملة', 'currency', 'SAR'),
            ('الموقع', 'location', 'الرياض، حي النرجس'),
            ('العميل', 'client', 'شركة التطوير العقاري'),
            ('مدير المشروع', 'manager', 'أحمد محمد'),
            ('نوع المشروع', 'type', 'إنشائي')
        ]
        
        table = doc.add_table(rows=len(project_fields)+1, cols=3)
        table.style = 'Table Grid'
        
        # رأس الجدول
        header_cells = table.rows[0].cells
        header_cells[0].text = 'الحقل'
        header_cells[1].text = 'القيمة'
        header_cells[2].text = 'مثال'
        
        # تعبئة البيانات
        for i, (field_name, field_key, example) in enumerate(project_fields, 1):
            cells = table.rows[i].cells
            cells[0].text = field_name
            if project_data and project_data.get(field_key):
                cells[1].text = str(project_data.get(field_key, ''))
            else:
                cells[1].text = '__________'
            cells[2].text = example
        
        doc.add_paragraph()
        
        # ============================================
        # المهام
        # ============================================
        doc.add_heading('مهام المشروع', level=1)
        
        tasks_table = doc.add_table(rows=1, cols=9)
        tasks_table.style = 'Table Grid'
        
        # رأس جدول المهام
        header_cells = tasks_table.rows[0].cells
        headers = ['الرقم', 'اسم المهمة', 'الوصف', 'المسؤول', 'تاريخ البدء', 
                   'تاريخ الانتهاء', 'المدة', 'الأولوية', 'المهام السابقة']
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        # صفوف المهام
        if project_data and project_data.get('tasks'):
            for i, task in enumerate(project_data['tasks'], 1):
                row = tasks_table.add_row()
                row.cells[0].text = str(i)
                row.cells[1].text = task.get('name', '')
                row.cells[2].text = task.get('description', '')
                row.cells[3].text = task.get('assigned_to', '')
                row.cells[4].text = task.get('start_date', '')
                row.cells[5].text = task.get('end_date', '')
                row.cells[6].text = str(task.get('duration', ''))
                row.cells[7].text = task.get('priority', 'متوسطة')
                row.cells[8].text = task.get('depends_on', '')
        else:
            # 10 صفوف فارغة
            for i in range(1, 11):
                row = tasks_table.add_row()
                row.cells[0].text = str(i)
                for j in range(1, 9):
                    row.cells[j].text = '__________'
        
        doc.add_paragraph()
        
        # ============================================
        # الموارد
        # ============================================
        doc.add_heading('الموارد المطلوبة', level=1)
        
        resources_table = doc.add_table(rows=4, cols=2)
        resources_table.style = 'Table Grid'
        
        resources_data = [
            ('المواد', 'خرسانة، حديد، بلوك، أسمنت'),
            ('المعدات', 'خلاطة خرسانة، رافعة، حفار'),
            ('المهارات', 'حدادة، نجارة، كهرباء'),
            ('ملاحظات', 'أي موارد إضافية')
        ]
        
        for i, (resource, example) in enumerate(resources_data):
            row = resources_table.rows[i]
            row.cells[0].text = resource
            row.cells[1].text = f'__________  (مثال: {example})'
        
        # حفظ المستند
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    
    # ============================================
    # توليد قالب PDF
    # ============================================
    
    def generate_pdf_template(self, project_data=None):
        """توليد قالب PDF مع دعم العربية"""
        output = io.BytesIO()
        doc = SimpleDocTemplate(output, pagesize=A4)
        elements = []
        
        # أنماط النصوص
        styles = getSampleStyleSheet()
        title_style = styles['Title']
        title_style.alignment = 1  # توسيط
        
        heading_style = styles['Heading2']
        heading_style.alignment = 2  # يمين
        
        normal_style = styles['Normal']
        normal_style.alignment = 2  # يمين
        normal_style.fontName = 'Arial'
        normal_style.fontSize = 11
        
        # عنوان المستند
        elements.append(Paragraph("نموذج مشروع", title_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # ============================================
        # معلومات المشروع
        # ============================================
        elements.append(Paragraph("معلومات المشروع", heading_style))
        elements.append(Spacer(1, 0.1*inch))
        
        project_data_table = [
            ['الحقل', 'القيمة', 'مثال'],
            ['اسم المشروع', '__________', 'مشروع إنشاء مبنى إداري'],
            ['رمز المشروع', '__________', 'PRJ-2024-001'],
            ['وصف المشروع', '__________', 'وصف مختصر للمشروع'],
            ['تاريخ البدء', '__________', '2024-01-15'],
            ['تاريخ الانتهاء', '__________', '2024-12-31'],
            ['الميزانية', '__________', '1,500,000'],
            ['العملة', '__________', 'SAR'],
            ['الموقع', '__________', 'الرياض، حي النرجس'],
            ['العميل', '__________', 'شركة التطوير العقاري'],
            ['مدير المشروع', '__________', 'أحمد محمد'],
            ['نوع المشروع', '__________', 'إنشائي']
        ]
        
        table = Table(project_data_table, colWidths=[1.5*inch, 2*inch, 2*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Arial'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(table)
        elements.append(Spacer(1, 0.2*inch))
        
        # ============================================
        # المهام
        # ============================================
        elements.append(Paragraph("مهام المشروع", heading_style))
        elements.append(Spacer(1, 0.1*inch))
        
        tasks_headers = ['#', 'اسم المهمة', 'المسؤول', 'المدة', 'الأولوية']
        tasks_data = [tasks_headers]
        
        for i in range(1, 11):
            tasks_data.append([
                str(i),
                '__________',
                '__________',
                '__________',
                '__________'
            ])
        
        tasks_table = Table(tasks_data, colWidths=[0.5*inch, 3*inch, 1.5*inch, 1*inch, 1.2*inch])
        tasks_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), 'Arial'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(tasks_table)
        elements.append(Spacer(1, 0.2*inch))
        
        # ============================================
        # الموارد
        # ============================================
        elements.append(Paragraph("الموارد المطلوبة", heading_style))
        elements.append(Spacer(1, 0.1*inch))
        
        resources_data = [
            ['نوع المورد', 'الوصف', 'مثال'],
            ['المواد', '__________', 'خرسانة، حديد، بلوك'],
            ['المعدات', '__________', 'خلاطة خرسانة، رافعة'],
            ['المهارات', '__________', 'حدادة، نجارة، كهرباء']
        ]
        
        resources_table = Table(resources_data, colWidths=[1.5*inch, 2.5*inch, 2*inch])
        resources_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Arial'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        elements.append(resources_table)
        
        # بناء PDF
        doc.build(elements)
        output.seek(0)
        return output
    
    # ============================================
    # توليد قالب CSV
    # ============================================
    
    def generate_csv_template(self, project_data=None):
        """توليد قالب CSV مع دعم العربية"""
        output = io.StringIO()
        writer = csv.writer(output)
        
        # كتابة BOM لدعم العربية في Excel
        output.write('\ufeff')
        
        # معلومات المشروع
        writer.writerow(['نوع البيانات', 'القيمة', 'ملاحظات'])
        writer.writerow(['اسم المشروع', '', 'أدخل اسم المشروع'])
        writer.writerow(['رمز المشروع', '', 'رمز فريد للمشروع'])
        writer.writerow(['الوصف', '', 'وصف المشروع'])
        writer.writerow(['تاريخ البدء', '', 'YYYY-MM-DD'])
        writer.writerow(['تاريخ الانتهاء', '', 'YYYY-MM-DD'])
        writer.writerow(['الميزانية', '', 'القيمة بالأرقام'])
        writer.writerow(['العملة', 'SAR', 'SAR, USD, EUR'])
        writer.writerow(['الموقع', '', 'عنوان المشروع'])
        writer.writerow(['العميل', '', 'اسم العميل'])
        writer.writerow(['مدير المشروع', '', 'اسم مدير المشروع'])
        
        writer.writerow([])
        writer.writerow(['المهام'])
        writer.writerow(['اسم المهمة', 'الوصف', 'المسؤول', 'تاريخ البدء', 'تاريخ الانتهاء', 'المدة', 'الأولوية', 'المهام السابقة'])
        
        # 10 صفوف فارغة للمهام
        for _ in range(10):
            writer.writerow(['', '', '', '', '', '', 'متوسطة', ''])
        
        output.seek(0)
        return io.BytesIO(output.getvalue().encode('utf-8-sig'))
    
    # ============================================
    # توليد قالب JSON
    # ============================================
    
    def generate_json_template(self, project_data=None):
        """توليد قالب JSON مع دعم العربية"""
        template = {
            "project": {
                "name": "",
                "code": "",
                "description": "",
                "start_date": "",
                "end_date": "",
                "budget": 0,
                "currency": "SAR",
                "location": "",
                "client": "",
                "manager": "",
                "type": ""
            },
            "tasks": [
                {
                    "id": i,
                    "name": "",
                    "description": "",
                    "assigned_to": "",
                    "start_date": "",
                    "end_date": "",
                    "duration": 1,
                    "priority": "medium",
                    "depends_on": [],
                    "resources": []
                } for i in range(1, 11)
            ],
            "resources": {
                "materials": [],
                "equipment": [],
                "skills": []
            },
            "metadata": {
                "version": "1.0",
                "language": "ar",
                "created_at": datetime.now().isoformat(),
                "description": "نموذج مشروع - املأ البيانات ثم ارفع الملف"
            }
        }
        
        if project_data:
            # دمج البيانات الموجودة
            if 'project' in project_data:
                template['project'].update(project_data['project'])
            if 'tasks' in project_data:
                for i, task_data in enumerate(project_data['tasks'][:10]):
                    template['tasks'][i].update(task_data)
        
        output = io.BytesIO()
        output.write(json.dumps(template, ensure_ascii=False, indent=2).encode('utf-8'))
        output.seek(0)
        return output
    
    # ============================================
    # الدالة الرئيسية لتوليد القالب
    # ============================================
    
    def generate_template(self, format_type='excel', project_data=None):
        """توليد قالب بالتنسيق المطلوب"""
        generators = {
            'excel': self.generate_excel_template,
            'word': self.generate_word_template,
            'pdf': self.generate_pdf_template,
            'csv': self.generate_csv_template,
            'json': self.generate_json_template
        }
        
        generator = generators.get(format_type.lower())
        if not generator:
            return None, 'نوع الملف غير مدعوم'
        
        try:
            output = generator(project_data)
            return output, None
        except Exception as e:
            return None, str(e)
    
    def get_mime_type(self, format_type):
        """الحصول على MIME type للتنسيق المطلوب"""
        mime_types = {
            'excel': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'word': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'pdf': 'application/pdf',
            'csv': 'text/csv',
            'json': 'application/json'
        }
        return mime_types.get(format_type.lower(), 'application/octet-stream')
    
    def get_filename(self, format_type, project_name='project'):
        """توليد اسم الملف"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        extensions = {
            'excel': 'xlsx',
            'word': 'docx',
            'pdf': 'pdf',
            'csv': 'csv',
            'json': 'json'
        }
        ext = extensions.get(format_type.lower(), 'txt')
        safe_name = ''.join(c for c in project_name if c.isalnum() or c in ' -_').strip()
        return f"{safe_name}_template_{timestamp}.{ext}"