"""
document_parser.py - خدمة تحليل المستندات (Excel, Word, PDF)
"""
import os
import pandas as pd
import PyPDF2
from docx import Document
import json
from datetime import datetime
from werkzeug.utils import secure_filename
import re

class DocumentParser:
    """محلل المستندات"""
    
    def __init__(self, upload_folder):
        self.upload_folder = upload_folder
    
    def parse_document(self, file_path, file_extension):
        """تحليل المستند حسب نوعه"""
        try:
            if file_extension.lower() in ['xlsx', 'xls']:
                return self.parse_excel(file_path)
            elif file_extension.lower() == 'pdf':
                return self.parse_pdf(file_path)
            elif file_extension.lower() in ['docx', 'doc']:
                return self.parse_word(file_path)
            elif file_extension.lower() == 'csv':
                return self.parse_csv(file_path)
            else:
                raise ValueError(f'نوع الملف غير مدعوم: {file_extension}')
        except Exception as e:
            raise Exception(f'خطأ في تحليل المستند: {str(e)}')
    
    def parse_excel(self, file_path):
        """تحليل ملف Excel"""
        try:
            # قراءة ملف Excel
            xls = pd.ExcelFile(file_path)
            
            # الحصول على أسماء الأوراق
            sheet_names = xls.sheet_names
            
            results = {
                'type': 'excel',
                'sheet_count': len(sheet_names),
                'sheets': []
            }
            
            for sheet_name in sheet_names:
                # قراءة كل ورقة
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # تحليل البيانات
                sheet_data = self.analyze_excel_sheet(df, sheet_name)
                results['sheets'].append(sheet_data)
            
            # استخراج معلومات المشروع
            project_info = self.extract_project_info(results)
            
            # استخراج جدول الكميات
            bill_of_quantities = self.extract_bill_of_quantities(results)
            
            return {
                'success': True,
                'metadata': results,
                'project_info': project_info,
                'bill_of_quantities': bill_of_quantities,
                'processed_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            raise Exception(f'خطأ في تحليل ملف Excel: {str(e)}')
    
    def analyze_excel_sheet(self, df, sheet_name):
        """تحليل ورقة Excel"""
        # تنظيف البيانات
        df = df.dropna(how='all').dropna(axis=1, how='all')
        
        # تحويل إلى قاموس
        records = df.to_dict('records')
        columns = df.columns.tolist()
        
        # البحث عن العناوين الهامة
        titles = self.find_titles_in_dataframe(df)
        
        return {
            'name': sheet_name,
            'row_count': len(df),
            'column_count': len(columns),
            'columns': columns,
            'titles': titles,
            'sample_data': records[:10] if records else []
        }
    
    def find_titles_in_dataframe(self, df):
        """البحث عن عناوين المشروع في البيانات"""
        titles = {}
        
        # البحث في الصفوف الأولى للعناوين
        for i in range(min(10, len(df))):
            row = df.iloc[i]
            for col in df.columns:
                cell_value = str(row[col]) if pd.notna(row[col]) else ''
                
                # البحث عن عناوين المشروع
                if any(keyword in cell_value.lower() for keyword in ['مشروع', 'project']):
                    titles['project'] = cell_value
                elif any(keyword in cell_value.lower() for keyword in ['موقع', 'site']):
                    titles['site'] = cell_value
                elif any(keyword in cell_value.lower() for keyword in ['منطقة', 'area']):
                    titles['area'] = cell_value
                elif any(keyword in cell_value.lower() for keyword in ['عقد', 'contract']):
                    titles['contract'] = cell_value
        
        return titles
    
    def extract_project_info(self, excel_data):
        """استخراج معلومات المشروع من البيانات"""
        project_info = {
            'name': '',
            'site_name': '',
            'area_name': '',
            'contract_number': ''
        }
        
        for sheet in excel_data.get('sheets', []):
            for key, value in sheet.get('titles', {}).items():
                if key in project_info and not project_info[key]:
                    project_info[key] = value
        
        return project_info
    
    def extract_bill_of_quantities(self, excel_data):
        """استخراج جدول الكميات والمواصفات"""
        bill_items = []
        
        for sheet in excel_data.get('sheets', []):
            sheet_name = sheet.get('name', '').lower()
            
            # البحث عن أوراق جدول الكميات
            if any(keyword in sheet_name for keyword in ['جدول', 'كميات', 'مواصفات', 'bill', 'quantity']):
                columns = sheet.get('columns', [])
                data = sheet.get('sample_data', [])
                
                # البحث عن الأعمدة المطلوبة
                column_mapping = self.map_boq_columns(columns)
                
                if column_mapping:
                    # استخراج البيانات
                    for row in data:
                        item = {}
                        for target_col, source_col in column_mapping.items():
                            if source_col in row:
                                item[target_col] = row[source_col]
                        
                        if any(item.values()):  # إذا كان هناك بيانات
                            bill_items.append(item)
        
        return {
            'total_items': len(bill_items),
            'columns_found': list(set([col for item in bill_items for col in item.keys()])),
            'items': bill_items[:50]  # أول 50 عنصر فقط
        }
    
    def map_boq_columns(self, columns):
        """تعيين أعمدة جدول الكميات"""
        column_mapping = {}
        
        # التعرف على الأعمدة العربية
        arabic_mappings = {
            'item_code': ['رمز', 'كود', 'رقم', 'رمز البند'],
            'description': ['بند', 'وصف', 'العمل', 'مواصفات', 'وصف الأعمال'],
            'unit': ['وحدة', 'الوحدة'],
            'quantity': ['كمية', 'الكمية'],
            'unit_price': ['سعر', 'سعر الوحدة', 'السعر'],
            'total_price': ['إجمالي', 'المجموع', 'القيمة'],
            'notes': ['ملاحظات', 'ملاحظة']
        }
        
        for col in columns:
            col_str = str(col).lower()
            
            for target, keywords in arabic_mappings.items():
                for keyword in keywords:
                    if keyword in col_str and target not in column_mapping:
                        column_mapping[target] = col
                        break
        
        return column_mapping
    
    def parse_pdf(self, file_path):
        """تحليل ملف PDF"""
        try:
            text_content = ""
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # استخراج النص من جميع الصفحات
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text()
            
            # تحليل النص
            project_info = self.extract_info_from_text(text_content)
            tables = self.extract_tables_from_text(text_content)
            
            return {
                'success': True,
                'type': 'pdf',
                'page_count': len(pdf_reader.pages),
                'text_length': len(text_content),
                'project_info': project_info,
                'tables_found': len(tables),
                'tables': tables[:5],  # أول 5 جداول فقط
                'processed_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            raise Exception(f'خطأ في تحليل ملف PDF: {str(e)}')
    
    def extract_info_from_text(self, text):
        """استخراج معلومات من النص"""
        info = {
            'project_name': '',
            'site_name': '',
            'area_name': '',
            'contract_number': ''
        }
        
        # أنماط البحث
        patterns = {
            'project_name': r'(?:مشروع|project)[:\s]+(.+)',
            'site_name': r'(?:موقع|site)[:\s]+(.+)',
            'area_name': r'(?:منطقة|area)[:\s]+(.+)',
            'contract_number': r'(?:عقد|رقم العقد|contract)[:\s]+(.+)'
        }
        
        for key, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                info[key] = match.group(1).strip()
        
        return info
    
    def extract_tables_from_text(self, text):
        """استخراج الجداول من النص"""
        tables = []
        
        # تقسيم النص إلى سطور
        lines = text.split('\n')
        
        current_table = []
        in_table = False
        
        for line in lines:
            # البحث عن بداية الجدول (سطور تحتوي على أعمدة)
            if re.search(r'[\d,.]+[\s]+[\d,.]+', line):  # خط يحتوي على أرقام
                if not in_table:
                    in_table = True
                current_table.append(line.strip())
            elif in_table and line.strip() == '':
                # نهاية الجدول
                if current_table:
                    tables.append(current_table.copy())
                    current_table = []
                in_table = False
        
        return tables
    
    def parse_word(self, file_path):
        """تحليل ملف Word"""
        try:
            doc = Document(file_path)
            
            # استخراج النص
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            text_content = '\n'.join(full_text)
            
            # استخراج الجداول
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            # تحليل النص
            project_info = self.extract_info_from_text(text_content)
            
            return {
                'success': True,
                'type': 'word',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'text_length': len(text_content),
                'project_info': project_info,
                'tables': tables_data[:10],  # أول 10 جداول فقط
                'processed_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            raise Exception(f'خطأ في تحليل ملف Word: {str(e)}')
    
    def parse_csv(self, file_path):
        """تحليل ملف CSV"""
        try:
            # قراءة ملف CSV
            df = pd.read_csv(file_path, encoding='utf-8')
            
            # تحليل البيانات
            results = {
                'type': 'csv',
                'row_count': len(df),
                'column_count': len(df.columns),
                'columns': df.columns.tolist(),
                'sample_data': df.head(20).to_dict('records')
            }
            
            # استخراج معلومات المشروع
            project_info = self.extract_project_info_from_csv(df)
            
            return {
                'success': True,
                'metadata': results,
                'project_info': project_info,
                'processed_at': datetime.utcnow().isoformat()
            }
            
        except Exception as e:
            raise Exception(f'خطأ في تحليل ملف CSV: {str(e)}')
    
    def extract_project_info_from_csv(self, df):
        """استخراج معلومات المشروع من CSV"""
        info = {
            'project_name': '',
            'site_name': '',
            'area_name': ''
        }
        
        # البحث في أسماء الأعمدة
        for col in df.columns:
            col_lower = str(col).lower()
            
            if 'مشروع' in col_lower or 'project' in col_lower:
                # الحصول على قيمة غير فارغة من العمود
                non_null_values = df[col].dropna()
                if not non_null_values.empty:
                    info['project_name'] = str(non_null_values.iloc[0])
            
            elif 'موقع' in col_lower or 'site' in col_lower:
                non_null_values = df[col].dropna()
                if not non_null_values.empty:
                    info['site_name'] = str(non_null_values.iloc[0])
            
            elif 'منطقة' in col_lower or 'area' in col_lower:
                non_null_values = df[col].dropna()
                if not non_null_values.empty:
                    info['area_name'] = str(non_null_values.iloc[0])
        
        return info
    
    def save_parsed_data(self, project_id, parsed_data):
        """حفظ البيانات المحللة في قاعدة البيانات"""
        try:
            from uploads.temp.models import ProjectDocument, BillItem
            
            # حفظ المستند
            document = ProjectDocument(
                project_id=project_id,
                document_type='bill_of_quantities',
                extraction_metadata=parsed_data.get('metadata', {}),
                analysis_summary={
                    'project_info': parsed_data.get('project_info', {}),
                    'bill_items_count': parsed_data.get('bill_of_quantities', {}).get('total_items', 0)
                },
                extraction_status='completed'
            )
            
            # TODO: إضافة المستند إلى قاعدة البيانات
            
            # حفظ بنود الجدول
            bill_items_data = parsed_data.get('bill_of_quantities', {}).get('items', [])
            for item_data in bill_items_data:
                bill_item = BillItem(
                    project_id=project_id,
                    item_code=item_data.get('item_code', ''),
                    description=item_data.get('description', ''),
                    unit=item_data.get('unit', ''),
                    planned_quantity=float(item_data.get('quantity', 0) or 0),
                    unit_price=float(item_data.get('unit_price', 0) or 0)
                )
                bill_item.calculate_amounts()
                
                # TODO: إضافة البند إلى قاعدة البيانات
            
            return True
            
        except Exception as e:
            print(f"خطأ في حفظ البيانات المحللة: {str(e)}")
            return False